import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from parser_v2 import ParsedCV, Section, ExperienceBlock, TableRow

BLUE = "005297"

MONTHS = {
    "jan": 1, "janv": 1, "janvier": 1,
    "fev": 2, "fév": 2, "fevr": 2, "févr": 2, "fevrier": 2, "février": 2,
    "mar": 3, "mars": 3,
    "avr": 4, "avril": 4,
    "mai": 5,
    "juin": 6,
    "jul": 7, "juil": 7, "juillet": 7,
    "aou": 8, "aoû": 8, "aout": 8, "août": 8,
    "sep": 9, "sept": 9, "septembre": 9,
    "oct": 10, "octobre": 10,
    "nov": 11, "novembre": 11,
    "dec": 12, "déc": 12, "decembre": 12, "décembre": 12,
}


def _clear_body_keep_sections(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill="FFFFFF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_table_borders(table, color="BFBFBF", size=6):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def _add_right_tab(paragraph, position=9044):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    right_tab = OxmlElement("w:tab")
    right_tab.set(qn("w:val"), "right")
    right_tab.set(qn("w:pos"), str(position))
    tabs.append(right_tab)


def _set_run(run, *, bold=None, italic=None, size=None, color=None, caps=None):
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if caps is not None:
        run.font.all_caps = caps


def _style_name(doc, preferred, fallback="Normal"):
    """Find a style in the document, trying French equivalents if English not found."""
    # French ↔ English style name equivalents
    _FR_EQUIVALENTS = {
        "Heading 1": "Titre 1",
        "Heading 2": "Titre 2",
        "Heading 3": "Titre 3",
        "Heading 4": "Titre 4",
        "Heading 5": "Titre 5",
        "List Bullet": "Liste à puces",
        "List Paragraph": "Paragraphe de liste",
        "Titre 1": "Heading 1",
        "Titre 2": "Heading 2",
        "Titre 3": "Heading 3",
        "Titre 4": "Heading 4",
        "Titre 5": "Heading 5",
        "Liste à puces": "List Bullet",
        "Paragraphe de liste": "List Paragraph",
    }
    # Try preferred name first
    try:
        doc.styles[preferred]
        return preferred
    except Exception:
        pass
    # Try the French/English equivalent
    equiv = _FR_EQUIVALENTS.get(preferred)
    if equiv:
        try:
            doc.styles[equiv]
            return equiv
        except Exception:
            pass
    # Final fallback
    return fallback


def _prompt_with_default(label: str, current_value: str = "") -> str:
    current_value = (current_value or "").strip()
    prompt = f"{label} [{current_value}]: " if current_value else f"{label}: "
    value = input(prompt).strip()
    return value or current_value


def _add_spacer(doc, count=3):
    """Add empty paragraphs between sections to match the template spacing."""
    spacer_style = _style_name(doc, "Profil : Experience", "Normal")
    for _ in range(count):
        p = doc.add_paragraph(style=spacer_style)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("")
        run.font.size = Pt(6)


def _normalize_month(token: str) -> str:
    return token.lower().replace(".", "").strip()


def _parse_date_piece(piece: str):
    piece = piece.strip().replace("–", "-").replace("—", "-")
    # MM/YYYY format
    m = re.match(r"^(\d{2})/(\d{4})$", piece)
    if m:
        return int(m.group(2)), int(m.group(1))
    # Month YYYY format
    m = re.match(r"(?i)^([A-Za-zÀ-ÿ\.]+)\s+(\d{4})$", piece)
    if m:
        month = MONTHS.get(_normalize_month(m.group(1)))
        year = int(m.group(2))
        if month:
            return year, month
    # Just YYYY
    m = re.match(r"^(\d{4})$", piece)
    if m:
        return int(m.group(1)), 1
    return None


def _duration_label(dates: str) -> str:
    cleaned = dates.strip()
    if not cleaned:
        return ""

    # Already has "Durée" → keep as-is
    if cleaned.lower().startswith("durée"):
        return cleaned

    if re.fullmatch(r"\d{4}", cleaned):
        return f"Durée {cleaned}"

    parts = re.split(r"\s*[-–—]\s*", cleaned)
    if len(parts) == 2:
        start = _parse_date_piece(parts[0])
        end = _parse_date_piece(parts[1])
        if start and end:
            sy, sm = start
            ey, em = end
            months = (ey - sy) * 12 + (em - sm)
            if months <= 0:
                return f"Durée {cleaned}"
            if months < 12:
                return f"Durée {months} mois"
            years = months // 12
            rem = months % 12
            if rem == 0:
                return "Durée 1 an" if years == 1 else f"Durée {years} ans"
            return f"Durée 1 an {rem} mois" if years == 1 else f"Durée {years} ans {rem} mois"

    return f"Durée {cleaned}"


def _render_profile(doc, profile, language_override=None, exp_override=None, ask_user=True):
    profile_style = _style_name(doc, "Profil", "Normal")
    spacer_style = _style_name(doc, "Profil : Experience", "Normal")

    display_lang = language_override or profile.language or ""
    display_exp = exp_override or profile.years_experience or ""

    if ask_user:
        display_lang = _prompt_with_default("Langue", display_lang)
        display_exp = _prompt_with_default("Années d'expérience", display_exp)

    if display_exp and "expérience" not in display_exp.lower():
        display_exp = f"{display_exp} d'expérience professionnelle"

    for text in [profile.name, profile.title, display_lang, display_exp]:
        if not text:
            continue
        p = doc.add_paragraph(style=profile_style)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        _set_run(run, bold=True, color=BLUE)

    doc.add_paragraph(style=spacer_style)


def _render_section_header(doc, text):
    p = doc.add_paragraph(style=_style_name(doc, "Heading 1", "Normal"))
    run = p.add_run(text.upper())
    _set_run(run, bold=True, color="FFFFFF", size=14)


def _render_table(doc, rows: list[TableRow], left_cm=7.8, right_cm=8.2):
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_borders(table, color="BFBFBF", size=6)

    for i, row in enumerate(rows):
        left_cell = table.cell(i, 0)
        right_cell = table.cell(i, 1)

        left_cell.width = Cm(left_cm)
        right_cell.width = Cm(right_cm)

        left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        _set_cell_margins(left_cell)
        _set_cell_margins(right_cell)
        _set_cell_shading(left_cell, "FFFFFF")
        _set_cell_shading(right_cell, "FFFFFF")

        left_cell.text = ""
        right_cell.text = ""

        lp = left_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(row.left or "")
        _set_run(lr, bold=True, size=11)

        rp = right_cell.paragraphs[0]
        rp.paragraph_format.space_before = Pt(0)
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(row.right or "")
        _set_run(rr, size=11)


def _render_bullet(doc, text):
    p = doc.add_paragraph(style=_style_name(doc, "List Bullet", "List Paragraph"))
    p.add_run(text)


def _render_bullets(doc, items: list[str]):
    for item in items:
        if item:  # skip empty items
            _render_bullet(doc, item)


def _parse_title_line(title_line: str) -> tuple[str, str]:
    # Split on tab
    if "\t" in title_line:
        parts = title_line.split("\t", 1)
        left = parts[0].strip()
        right = parts[1].strip()

        # Check if right side is a date range → convert to duration
        if re.search(r"\d{2}/\d{4}\s*[-–—]\s*\d{2}/\d{4}", right):
            right = _duration_label(right)
        elif re.match(r"(?i)^durée", right):
            pass  # already a duration
        elif re.match(r"(?i)^depuis", right):
            pass  # open-ended duration like "Depuis Mars 2022"
        else:
            # Right side is a subtitle (e.g. project description) → merge with left
            left = f"{left} {right}"
            right = ""

        return left, right

    return title_line, ""


def _render_experience_block(doc, exp: ExperienceBlock):
    title_style = _style_name(doc, "Titre Référence", "Heading 2")
    poste_style = _style_name(doc, "Heading 3", "Normal")
    sub_style = _style_name(doc, "Heading 4", "Normal")

    # --- Title line with right-aligned duration ---
    left_text, right_text = _parse_title_line(exp.title_line)

    p = doc.add_paragraph(style=title_style)
    _add_right_tab(p)

    left_run = p.add_run(left_text)
    _set_run(left_run, bold=True, size=16, color=BLUE, caps=True)

    if right_text:
        right_run = p.add_run("\t" + right_text)
        _set_run(right_run, bold=True, size=16, color=BLUE, caps=True)

    # --- Poste line ---
    if exp.poste:
        p = doc.add_paragraph(style=poste_style)
        r = p.add_run(exp.poste)
        _set_run(r, italic=True, color=BLUE, size=12)

    # --- Sub-sections (ROLE, Environnement technique, Contexte, etc.) ---
    for sub_header, items in exp.sub_sections:
        if sub_header:
            p = doc.add_paragraph(style=sub_style)
            r = p.add_run(sub_header)
            _set_run(r, bold=True, color=BLUE, caps=True)

        for item in items:
            if item:
                _render_bullet(doc, item)


def _render_section(doc, section: Section):
    _render_section_header(doc, section.header)

    # Render table if present
    if section.table_rows:
        _render_table(doc, section.table_rows)

    # Render bullets if present
    if section.bullet_items:
        _render_bullets(doc, section.bullet_items)

    # Render experience blocks if present
    if section.experience_blocks:
        for exp in section.experience_blocks:
            _render_experience_block(doc, exp)

    # Add spacing after every section
    _add_spacer(doc, 3)


def generate_docx(
    cv: ParsedCV,
    output_path: str,
    template_path: str,
    language: str | None = None,
    years_experience: str | None = None,
    ask_user: bool = True,
):
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    _clear_body_keep_sections(doc)

    # --- Profile ---
    _render_profile(doc, cv.profile, language_override=language,
                    exp_override=years_experience, ask_user=ask_user)

    # --- Sections (in order, as found in the input) ---
    for section in cv.sections:
        _render_section(doc, section)

    doc.save(output_path)
    print(f"✅ Generated: {output_path}")