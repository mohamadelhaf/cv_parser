import re
from dataclasses import dataclass, field
from docx import Document


@dataclass
class ProfileData:
    lines: list[str] = field(default_factory=list)
    # Convenience accessors (filled after parsing)
    name: str = ""
    title: str = ""
    language: str = ""
    years_experience: str = ""


@dataclass
class TableRow:
    left: str = ""
    right: str = ""


@dataclass
class ExperienceBlock:
    title_line: str = ""          # From 'Titre Référence' style
    poste: str = ""               # From 'Heading 2' or 'Heading 3'
    sub_sections: list[tuple[str, list[str]]] = field(default_factory=list)
    # Each sub_section is (header_text, [bullet_items])
    # e.g. ("ROLE :", ["task1", "task2"]), ("Environnement technique :", ["tech1"])
    # or   ("Contexte", ["paragraph"]), ("Activités :", ["item1", "item2"])


@dataclass
class Section:
    header: str = ""              # Original header text (e.g. "SAVOIR FAIRE")
    content_type: str = ""        # "table", "bullets", "experiences", "empty"
    table_rows: list[TableRow] = field(default_factory=list)
    bullet_items: list[str] = field(default_factory=list)
    experience_blocks: list[ExperienceBlock] = field(default_factory=list)


@dataclass
class ParsedCV:
    profile: ProfileData = field(default_factory=ProfileData)
    sections: list[Section] = field(default_factory=list)


def parse_docx(path: str) -> ParsedCV:
    doc = Document(path)
    result = ParsedCV()
    
    # --- Pass 1: Read profile (Profil style) ---
    for para in doc.paragraphs:
        if para.style and para.style.name == "Profil":
            text = para.text.strip()
            if text:
                result.profile.lines.append(text)
    
    # Fill convenience fields
    if len(result.profile.lines) >= 1:
        result.profile.name = result.profile.lines[0]
    if len(result.profile.lines) >= 2:
        result.profile.title = result.profile.lines[1]
    if len(result.profile.lines) >= 3:
        result.profile.language = result.profile.lines[2]
    if len(result.profile.lines) >= 4:
        result.profile.years_experience = result.profile.lines[3]
    
    # --- Pass 2: Read sections by iterating body elements in order ---
    current_section: Section | None = None
    current_exp: ExperienceBlock | None = None
    current_sub_header: str = ""
    current_sub_items: list[str] = []
    
    body = doc.element.body
    
    for element in body:
        tag = element.tag.split("}")[-1]
        
        # --- Table element ---
        if tag == "tbl":
            if current_section is not None:
                # Flush any pending experience
                _flush_experience(current_section, current_exp, current_sub_header, current_sub_items)
                current_exp = None
                current_sub_header = ""
                current_sub_items = []
                
                # Read table rows
                current_section.content_type = "table"
                for tr in element.iter():
                    tr_tag = tr.tag.split("}")[-1]
                    if tr_tag == "tr":
                        cells = []
                        for tc in tr.iter():
                            tc_tag = tc.tag.split("}")[-1]
                            if tc_tag == "tc":
                                cell_texts = []
                                for node in tc.iter():
                                    node_tag = node.tag.split("}")[-1]
                                    if node_tag == "t" and node.text:
                                        cell_texts.append(node.text)
                                cells.append(" ".join(cell_texts).strip())
                        if len(cells) >= 2 and any(cells):
                            current_section.table_rows.append(
                                TableRow(left=cells[0], right=cells[1])
                            )
                        elif len(cells) == 1 and cells[0]:
                            current_section.table_rows.append(
                                TableRow(left=cells[0], right="")
                            )
            continue
        
        # --- Paragraph element ---
        if tag != "p":
            continue
        
        # Find the matching Paragraph object to get style
        para = None
        for p in doc.paragraphs:
            if p._p is element:
                para = p
                break
        
        if para is None:
            continue
        
        style = para.style.name if para.style else "Normal"
        
        # --- Normalize French style names to English equivalents ---
        _STYLE_MAP = {
            "Titre 1": "Heading 1",
            "Titre 2": "Heading 2",
            "Titre 3": "Heading 3",
            "Titre 4": "Heading 4",
            "Titre 5": "Heading 5",
            "Liste à puces": "List Bullet",
            "Paragraphe de liste": "List Paragraph",
            "Titre Référence": "Titre Référence",  # keep as-is
        }
        style = _STYLE_MAP.get(style, style)
        text = para.text.strip().replace("\t\n", " ").replace("\n", " ").replace("\xa0", " ").strip()
        
        # Skip empty paragraphs and profile (already handled)
        if style == "Profil" or style == "Profil : Experience":
            continue
        
        # --- Heading 1: new section ---
        if style == "Heading 1":
            # Flush previous section's pending experience
            if current_section is not None:
                _flush_experience(current_section, current_exp, current_sub_header, current_sub_items)
            
            # Start new section
            current_section = Section(header=text)
            result.sections.append(current_section)
            current_exp = None
            current_sub_header = ""
            current_sub_items = []
            continue
        
        # Everything below needs a current section
        if current_section is None:
            continue
        
        # --- Titre Référence: new experience/project entry ---
        if style == "Titre Référence":
            # Flush previous experience
            _flush_experience(current_section, current_exp, current_sub_header, current_sub_items)
            
            # Add "experiences" to content type (might already have "table")
            if current_section.content_type == "table":
                current_section.content_type = "table+experiences"
            else:
                current_section.content_type = "experiences"
            current_exp = ExperienceBlock(title_line=text)
            current_sub_header = ""
            current_sub_items = []
            continue
        
        # --- Heading 2 or Heading 3: poste/role name ---
        if style in ("Heading 2", "Heading 3"):
            if current_exp is not None:
                current_exp.poste = text
            continue
        
        # --- Heading 4: sub-header (ROLE, Environnement technique, Contexte, etc.) ---
        if style == "Heading 4":
            if current_exp is not None:
                # Save previous sub-section
                if current_sub_header or current_sub_items:
                    current_exp.sub_sections.append((current_sub_header, current_sub_items))
                current_sub_header = text
                current_sub_items = []
            continue
        
        # --- List Bullet: content item inside experience or standalone ---
        if style == "List Bullet":
            if current_exp is not None:
                current_sub_items.append(text)
            elif current_section.content_type != "table":
                current_section.content_type = "bullets"
                current_section.bullet_items.append(text)
            continue
        
        # --- List Paragraph: bullet-style items (Savoir Faire) ---
        if style == "List Paragraph":
            if current_exp is not None:
                current_sub_items.append(text)
            elif current_section.content_type != "table":
                current_section.content_type = "bullets"
                current_section.bullet_items.append(text)
            continue
        
        # --- Normal or other: detect experience titles in French INTM variant ---
        # French INTM variant uses Normal style with bold+tab for experience titles
        # e.g. "SNCF\t\t42 mois (2022 – 2026)" or "365Talents\t13 mois (2021 – 2022)"
        if style == "Normal" and text:
            # Check if this looks like an experience title line (has tab + duration pattern)
            has_tab = "\t" in para.text
            is_bold = all(r.bold for r in para.runs if r.text.strip()) if para.runs else False
            has_duration = bool(re.search(r"\d+\s*(?:mois|ans?|année)", text, re.IGNORECASE))
            
            if has_tab and is_bold and has_duration and current_section is not None:
                # This is an experience title line — treat like "Titre Référence"
                _flush_experience(current_section, current_exp, current_sub_header, current_sub_items)
                if current_section.content_type == "table":
                    current_section.content_type = "table+experiences"
                else:
                    current_section.content_type = "experiences"
                # Normalize tabs: replace multiple tabs with single tab
                normalized = re.sub(r"\t+", "\t", text)
                current_exp = ExperienceBlock(title_line=normalized)
                current_sub_header = ""
                current_sub_items = []
                continue
            
            # Check if this looks like a poste/role line (bold, short, no tab, inside experience)
            if is_bold and not has_tab and current_exp is not None and not current_exp.poste:
                word_count = len(text.split())
                if word_count <= 8 and not text.endswith(":"):
                    current_exp.poste = text
                    continue
        
        if text and current_exp is not None:
            # Check if it's an inline "Environnement Technique : ..." line or sub-header
            if current_sub_header or current_sub_items:
                current_exp.sub_sections.append((current_sub_header, current_sub_items))
            current_sub_header = text
            current_sub_items = []
        elif text and current_section is not None:
            # Standalone text in a section
            if current_section.content_type != "table":
                current_section.content_type = "bullets"
                current_section.bullet_items.append(text)
    
    # Flush last pending experience
    if current_section is not None:
        _flush_experience(current_section, current_exp, current_sub_header, current_sub_items)
    
    return result


def _flush_experience(section, exp, sub_header, sub_items):
    if exp is not None:
        if sub_header or sub_items:
            exp.sub_sections.append((sub_header, sub_items))
        section.experience_blocks.append(exp)
        # Only set content_type if not already set to table+experiences
        if "table" in section.content_type and "experiences" not in section.content_type:
            section.content_type = "table+experiences"
        elif section.content_type not in ("table+experiences",):
            section.content_type = "experiences"


def print_parsed(cv: ParsedCV):
    p = cv.profile
    print(f"\n{'='*60}")
    print(f"  PROFILE")
    print(f"    Name:       {p.name}")
    print(f"    Title:      {p.title}")
    print(f"    Language:   {p.language}")
    print(f"    Experience: {p.years_experience}")
    
    for i, section in enumerate(cv.sections):
        print(f"\n  SECTION {i+1}: [{section.header}] (type: {section.content_type})")
        
        if section.content_type == "table":
            for row in section.table_rows:
                print(f"    | {row.left[:30]:30s} | {row.right[:50]}")
        
        elif section.content_type == "bullets":
            for item in section.bullet_items:
                print(f"    • {item[:80]}")
        
        elif section.content_type == "experiences":
            for exp in section.experience_blocks:
                print(f"    [{exp.title_line[:60]}]")
                if exp.poste:
                    print(f"      Poste: {exp.poste}")
                for sub_header, items in exp.sub_sections:
                    print(f"      >> {sub_header}")
                    for item in items[:3]:
                        print(f"         • {item[:70]}")
                    if len(items) > 3:
                        print(f"         ... +{len(items)-3} more")
    
    print(f"\n{'='*60}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python parser_v2.py <docx_path>")
        sys.exit(1)
    
    cv = parse_docx(sys.argv[1])
    print_parsed(cv)